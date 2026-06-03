"""
Gera documentacao Word do Stock Tracker em word/StockTracker_Documentacao_Projeto.docx

Uso:
  pip install python-docx
  python tools/build_project_docx.py
"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "word"
OUT_FILE = OUT_DIR / "StockTracker_Documentacao_Projeto.docx"


def _heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True


def _bullet(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
    doc.add_paragraph()


def _add_evolution_section(doc: Document) -> None:
    """Historico das versoes desde consola sem GUI ate versao atual."""
    _heading(doc, "2. Evolucao do projeto e versoes", 1)
    _para(
        doc,
        "O Stock Tracker foi desenvolvido em fases incrementais. A logica de negocio "
        "(core/stock.py) manteve-se separada da interface desde cedo; as versoes "
        "reflectem sobretudo mudancas na GUI, no numero de distribuidores e na "
        "experiencia do utilizador. As versoes intermedias de demonstracao "
        "(run-demo.bat) foram unificadas numa unica aplicacao (run.bat).",
    )

    _heading(doc, "2.1 Linha do tempo (resumo)", 2)
    _table(
        doc,
        ["Fase", "Versao", "Interface", "Distribuidores", "Como arrancar"],
        [
            [
                "0 — Prototipo",
                "v0.1 Consola",
                "Nenhuma (menu texto)",
                "Excel; Mouser via codigo",
                "python -m src.test_terminal",
            ],
            [
                "1 — GUI inicial",
                "v0.2 Legada",
                "PySide6 simples (verde/vermelho)",
                "Apenas Mouser",
                "run-demo-legacy.bat (retirado)",
            ],
            [
                "2 — Siemens",
                "v1.0 Demo combo",
                "Template Siemens",
                "Um distribuidor escolhido",
                "run-demo.bat (retirado)",
            ],
            [
                "3 — Siemens+",
                "v1.1 Demo 2.0",
                "Siemens + popup nome",
                "Combo ou teste multi",
                "run-demo-2.bat (retirado)",
            ],
            [
                "4 — Producao",
                "v2.0 Atual",
                "Siemens completo (Qt Designer)",
                "Todos automaticos no SCAN",
                "run.bat / python -m src.main",
            ],
        ],
    )

    _heading(doc, "2.2 Fase 0 — Sem interface grafica (v0.1 Consola)", 2)
    _para(doc, "Primeira fase funcional: apenas Python + Excel + menu na consola.", bold=True)
    _bullet(
        doc,
        [
            "Classe StockTracker em src/core/stock.py criada desde o inicio.",
            "Ficheiro data/stock.xlsx como base de dados.",
            "Menu interativo: src/test_terminal.py (ainda disponivel para testes).",
            "Sem PySide6 — validacao rapida da logica antes do investimento em UI.",
        ],
    )
    _heading(doc, "Funcoes disponiveis na consola (v0.1)", 3)
    _table(
        doc,
        ["Opcao", "Funcao"],
        [
            ["1", "Pesquisar componente no Excel"],
            ["2", "Ver componente por codigo (simula scan)"],
            ["3", "Adicionar stock (IN)"],
            ["4", "Remover stock (OUT) com confirmacao"],
            ["5", "Pesquisar todos os catalogos configurados (consulta)"],
            ["6", "Catalogo → Excel → adicionar stock"],
            ["7", "Listar componentes no Excel"],
            ["8", "Ultimos 20 movimentos de historico"],
            ["9", "Alterar nome de utilizador"],
            ["10", "Adicionar componente manual"],
        ],
    )

    _heading(doc, "2.3 Fase 1 — Interface grafica legada (v0.2)", 2)
    _para(
        doc,
        "Primeira janela grafica: layout simples, botoes verdes (ADD STOCK) e "
        "vermelhos (REMOVE STOCK), sem template corporativo Siemens. Integracao "
        "apenas com API Mouser no fluxo SCAN.",
        bold=True,
    )
    _bullet(
        doc,
        [
            "Codigo em src/gui/demo/ (DemoStockTrackerWindow).",
            "Arranque historico: run-demo-legacy.bat → src/main_demo_legacy.py.",
            "Mensagens com QMessageBox padrao (antes dos popups Siemens).",
            "Funcoes: pesquisa Excel, scan Mouser, stock IN/OUT, historico basico.",
        ],
    )
    _para(
        doc,
        "Estado atual: pasta demo mantida no repositorio como referencia; "
        "nao faz parte do arranque oficial.",
    )

    _heading(doc, "2.4 Fase 2 — Template Siemens e escolha de distribuidor (v1.0)", 2)
    _para(
        doc,
        "Migracao para o design Siemens (cores #000028, acentos ciano, tipografia). "
        "Introducao do combo Distributor: o utilizador escolhia Mouser, DigiKey, etc. "
        "antes do scan.",
        bold=True,
    )
    _bullet(
        doc,
        [
            "Janela: src/gui/stock_tracker_window.py + designer/gui_stocktracker.ui.",
            "Arranque historico: run-demo.bat (main_demo.py, multi_catalog=False).",
            "Modulos suppliers/ para Mouser, DigiKey, TME, RS.",
            "Consola e GUI em paralelo para testes.",
        ],
    )

    _heading(doc, "2.5 Fase 3 — Demo 2.0 (refinamentos UX)", 2)
    _para(
        doc,
        "Iteracao sobre a GUI Siemens: popup para introduzir nome de utilizador "
        "quando em falta; experimentos com pesquisa automatica em todos os "
        "catalogos vs. combo de distribuidor.",
        bold=True,
    )
    _bullet(
        doc,
        [
            "user_name_dialog.py — pedir nome no proprio dialogo Siemens.",
            "Arranque historico: run-demo-2.bat.",
            "Testou-se multi_catalog=True (pesquisa em todos os APIs) e depois "
            "voltou-se ao combo para comparacao com v1.0.",
        ],
    )

    _heading(doc, "2.6 Fase 4 — Versao atual unificada (v2.0)", 2)
    _para(
        doc,
        "Versao de entrega: uma unica aplicacao, sem variantes demo no arranque.",
        bold=True,
    )
    _bullet(
        doc,
        [
            "Arranque unico: run.bat → python -m src.main.",
            "Sem combo Distributor — SCAN pesquisa automaticamente todos os APIs "
            "configurados (ordem: Mouser, TME, RS, DigiKey).",
            "Popups Siemens para confirmacao, historico, pesquisa Excel, manual, edicao.",
            "Botoes Copy nos campos de detalhes; EDIT COMPONENT; MANUAL COMPONENT.",
            "Qt Designer como fonte de verdade da UI (export via tools/export-ui.bat).",
            "Removidos: run-demo.bat, run-demo-2.bat, run-demo-legacy.bat.",
        ],
    )

    _heading(doc, "2.7 Comparativo de funcionalidades por versao", 2)
    _table(
        doc,
        ["Funcionalidade", "v0.1 Consola", "v0.2 Legada", "v1.0 Combo", "v2.0 Atual"],
        [
            ["Excel Components + History", "Sim", "Sim", "Sim", "Sim"],
            ["Interface grafica", "Nao", "Simples", "Siemens", "Siemens + Designer"],
            ["Pesquisa Excel", "Sim", "Sim", "Sim", "Sim + dialogo multi-resultado"],
            ["Scan + API distribuidor", "Sim (1)", "So Mouser", "Um escolhido", "Todos auto."],
            ["Stock IN / OUT", "Sim", "Sim", "Sim", "Sim"],
            ["Confirmacao OUT", "Sim", "Sim", "Sim", "Popup Siemens"],
            ["Historico", "Sim", "Sim", "Sim", "Popup tabela Siemens"],
            ["Componente manual", "Sim", "Nao", "Nao", "Popup dedicado"],
            ["Editar componente", "Nao", "Nao", "Nao", "Sim"],
            ["Autocomplete Excel", "Nao", "Nao", "Parcial", "Sim"],
            ["Copy para clipboard", "Nao", "Nao", "Nao", "Sim"],
            ["Popup nome utilizador", "Nao", "Nao", "v1.1 demo", "Sim"],
            ["DigiKey / TME / RS", "Core", "Nao", "Opcional", "Sim (config)"],
        ],
    )
    _para(
        doc,
        "(1) Na consola, escolha manual de fornecedor nas opcoes 5 e 6; na v2.0 o SCAN "
        "percorre todos os configurados.",
    )

    _heading(doc, "2.8 Projeto legado (referencia externa)", 2)
    _para(
        doc,
        "Existe um caminho antigo Documents\\stock-tracker — projeto separado. "
        "Nao misturar com este repositorio (StockTracker em Downloads).",
    )

    doc.add_page_break()


def build() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # Capa
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Stock Tracker\n")
    r.bold = True
    r.font.size = Pt(22)
    sub = title.add_run(
        "Documentacao do projeto\n"
        "Inventario de componentes eletronicos\n\n"
        "Siemens — projeto de estagio\n"
    )
    sub.font.size = Pt(14)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Versao 1.1 | {date.today().strftime('%d/%m/%Y')}\n")
    meta.add_run(f"Repositorio: {ROOT.name}")
    doc.add_page_break()

    _heading(doc, "Indice", 1)
    _para(
        doc,
        "No Microsoft Word: Referencias → Indice → Atualizar indice "
        "(utiliza os titulos deste documento).",
    )
    doc.add_page_break()

    _heading(doc, "1. Resumo executivo", 1)
    _para(
        doc,
        "O Stock Tracker e uma aplicacao desktop para Windows que gere inventario de "
        "componentes eletronicos. Os dados persistem em ficheiro Excel (stock.xlsx); "
        "a interface segue o design Siemens (PySide6 / Qt). O utilizador regista "
        "entradas e saidas de stock com historico auditavel, pesquisa componentes no "
        "Excel e, quando necessario, importa dados de catalogos de distribuidores "
        "(Mouser, TME, RS, DigiKey) atraves de APIs configuraveis. "
        "A evolucao desde a fase apenas-consola ate a versao atual esta na seccao 2.",
    )

    _add_evolution_section(doc)

    _heading(doc, "3. Contexto e objetivos", 1)
    _heading(doc, "3.1 Problema", 2)
    _bullet(
        doc,
        [
            "Gestao manual de referencias de fornecedor, fabricante e quantidades.",
            "Risco de erro ao atualizar stock sem registo de quem fez o movimento.",
            "Necessidade de consultar catalogos online ao receber novas pecas.",
        ],
    )
    _heading(doc, "3.2 Objetivos", 2)
    _bullet(
        doc,
        [
            "Centralizar inventario em Excel com folhas Components e History.",
            "Fornecer interface grafica alinhada aos padroes visuais Siemens.",
            "Automatizar consulta a APIs de distribuidores no fluxo de scan.",
            "Separar logica de negocio (core) da interface (gui) para manutencao.",
        ],
    )

    _heading(doc, "4. Requisitos funcionais", 1)
    _table(
        doc,
        ["ID", "Requisito", "Estado"],
        [
            ["RF01", "Identificacao do utilizador antes de operacoes", "Implementado"],
            ["RF02", "Pesquisa de componentes no Excel (SEARCH)", "Implementado"],
            ["RF03", "Scan de referencia / codigo de barras (SCAN)", "Implementado"],
            ["RF04", "Importacao via APIs se nao existir no Excel", "Implementado"],
            ["RF05", "Entrada e saida de stock (IN/OUT) com confirmacao OUT", "Implementado"],
            ["RF06", "Historico de movimentos (ultimos 20 / por componente)", "Implementado"],
            ["RF07", "Adicionar componente manual (sem API)", "Implementado"],
            ["RF08", "Editar dados de componente existente", "Implementado"],
            ["RF09", "Autocomplete a partir do Excel", "Implementado"],
            ["RF10", "Copiar campos para area de transferencia", "Implementado"],
        ],
    )

    _heading(doc, "5. Arquitetura do sistema", 1)
    _heading(doc, "5.1 Principio de separacao", 2)
    _para(doc, "Interface (gui/)  →  Logica (core/stock.py)  →  Dados (Excel / APIs)")
    _para(
        doc,
        "A camada grafica nao implementa regras de negocio: valida entrada do "
        "utilizador, mostra mensagens e invoca metodos da classe StockTracker.",
    )

    _heading(doc, "5.2 Estrutura de pastas", 2)
    _table(
        doc,
        ["Pasta / ficheiro", "Funcao"],
        [
            ["run.bat", "Arranque da aplicacao no Windows"],
            ["src/main.py", "Ponto de entrada da GUI"],
            ["src/core/stock.py", "Classe StockTracker — Excel, stock, historico"],
            ["src/core/suppliers/", "Integracao Mouser, DigiKey, TME, RS"],
            ["src/gui/stock_tracker_window.py", "Janela principal e eventos"],
            ["src/gui/designer/", "Ficheiros Qt Designer (.ui e export .py)"],
            ["config/secrets.py", "Credenciais API (local, nao versionar)"],
            ["data/stock.xlsx", "Base de dados Excel"],
            ["docs/", "Documentacao Markdown de apoio"],
            ["tools/", "Scripts de manutencao (Designer, export UI)"],
            ["word/", "Este documento Word"],
        ],
    )

    _heading(doc, "5.3 Classe StockTracker (negocio)", 2)
    _table(
        doc,
        ["Area", "Metodos principais"],
        [
            ["Excel", "get_workbook, get_components_sheet, save_workbook"],
            ["Pesquisa", "search_in_excel_all, find_component_any, extract_part_number"],
            ["Stock", "update_stock (IN/OUT), add_history"],
            ["Componentes", "add_component_row, add_manual_component, update_component"],
            ["Distribuidores", "search_any_supplier, search_suppliers_order"],
        ],
    )

    _heading(doc, "5.4 Fluxo SCAN (catalogo + Excel)", 2)
    _bullet(
        doc,
        [
            "Utilizador introduz nome (campo ou popup se vazio).",
            "Referencia no campo Scan Barcode / Supplier Ref. (minimo 5 caracteres).",
            "Procura no Excel; se encontrar, mostra detalhes.",
            "Se nao encontrar: confirma pesquisa nos catalogos configurados.",
            "Pesquisa automatica por ordem: Mouser → TME → RS → DigiKey → …",
            "Nova linha no Excel com stock 0; utilizador indica quantidade e ADD STOCK.",
        ],
    )

    _heading(doc, "6. Interface grafica", 1)
    _heading(doc, "6.1 Tecnologia", 2)
    _bullet(
        doc,
        [
            "PySide6 (Qt 6) — framework desktop multiplataforma.",
            "Layout principal: gui_stocktracker.ui (Qt Designer) exportado para Python.",
            "Popups Siemens: confirmacao, historico, pesquisa Excel, manual, edicao.",
            "Estilos corporativos em src/gui/styles.py e recursos em siemens_template/.",
        ],
    )

    _heading(doc, "6.2 Ecra principal — operacoes", 2)
    _table(
        doc,
        ["Campo / botao", "Funcao"],
        [
            ["User Name", "Identificacao para historico"],
            ["Search Component + SEARCH", "Pesquisa apenas no Excel"],
            ["Scan Barcode / Supplier Ref. + SCAN", "Excel + APIs de distribuidores"],
            ["Quantity + ADD / REMOVE STOCK", "Movimentos IN e OUT"],
            ["MANUAL COMPONENT", "Inserir sem API"],
            ["EDIT COMPONENT", "Alterar linha no Excel"],
            ["Historico", "Ultimos 20 ou por componente"],
            ["CLEAR / Exit", "Limpar campos / fechar"],
        ],
    )

    _heading(doc, "6.3 Regras na GUI", 2)
    _bullet(
        doc,
        [
            "Nome de utilizador obrigatorio (popup dedicado se em falta).",
            "Confirmacao explicita antes de remocao de stock (OUT).",
            "Referencia minima de 5 caracteres no scan.",
            "Mensagem critica se stock.xlsx estiver aberto no Microsoft Excel.",
        ],
    )

    _heading(doc, "7. Integracao com distribuidores", 1)
    _table(
        doc,
        ["Fornecedor", "Credenciais (secrets.py)", "Estado"],
        [
            ["Mouser", "MOUSER_API_KEY", "Operacional — pesquisa por part number"],
            ["TME", "TME_API_TOKEN, TME_APP_SECRET", "Implementado"],
            ["RS Components", "RS_API_KEY", "Implementado (URL em rs.py)"],
            [
                "DigiKey",
                "DIGIKEY_CLIENT_ID, DIGIKEY_CLIENT_SECRET, DIGIKEY_ENV",
                "OAuth OK; pesquisa sandbox pode devolver 403 (portal)",
            ],
            ["Robert Mauser", "ROBERT_MAUSER_API_KEY", "Reservado — sem API publica"],
        ],
    )
    _para(
        doc,
        "Ordem de pesquisa no SCAN: Mouser, TME, RS, DigiKey, Robert Mauser (apenas "
        "fornecedores com chaves preenchidas em config/secrets.py).",
    )
    _heading(doc, "7.1 DigiKey (nota tecnica)", 2)
    _para(
        doc,
        "A integracao usa OAuth 2-legged e Product Information V4. O token sandbox "
        "pode ser obtido com sucesso (HTTP 200) enquanto a pesquisa devolve 403 se "
        "a app no developer.digikey.com nao tiver a API totalmente autorizada. "
        "Documentacao detalhada: docs/user/DIGIKEY_SETUP.md.",
    )

    _heading(doc, "8. Modelo de dados (Excel)", 1)
    _heading(doc, "8.1 Folha Components", 2)
    _para(
        doc,
        "Colunas: ID, Supplier Reference, Manufacturer, Manufacturer Reference, "
        "Value, Description, Stock.",
    )
    _heading(doc, "8.2 Folha History", 2)
    _para(
        doc,
        "Colunas: Date, User, Supplier Reference, Movement (IN/OUT), Quantity, "
        "Stock After.",
    )

    _heading(doc, "9. Instalacao e execucao", 1)
    _heading(doc, "9.1 Requisitos", 2)
    _bullet(
        doc,
        [
            "Windows 10 ou 11",
            "Python 3.10 ou superior",
            "Ligacao a Internet (APIs de distribuidores)",
            "Microsoft Excel instalado (ficheiro fechado durante gravacao)",
        ],
    )
    _heading(doc, "9.2 Instalacao", 2)
    _bullet(
        doc,
        [
            "cd para a pasta do projeto StockTracker",
            "python -m venv .venv",
            ".venv\\Scripts\\activate",
            "pip install -r requirements.txt",
            "copy config\\secrets.example.py config\\secrets.py",
            "Editar secrets.py com chaves API (nunca commitar)",
            "Colocar ou criar data\\stock.xlsx",
        ],
    )
    _heading(doc, "9.3 Executar", 2)
    _table(
        doc,
        ["Metodo", "Comando"],
        [
            ["Duplo clique", "run.bat"],
            ["Terminal", "python -m src.main"],
            ["Testes (opcional)", "python -m src.test_terminal"],
            ["Teste DigiKey", "python scripts/test_digikey_auth.py"],
        ],
    )

    _heading(doc, "10. Qt Designer e manutencao da UI", 1)
    _bullet(
        doc,
        [
            "Editar layout: src/gui/designer/gui_stocktracker.ui (Qt Designer).",
            "Apos editar: tools\\export-ui.bat (exporta .ui → .py e corrige resources_rc).",
            "Nao editar gui_stocktracker.py manualmente — ficheiro gerado.",
            "Popups: designer/popups/*.ui — gerar com python tools/generate_popup_uis.py.",
            "Regenerar .ui base: python tools/generate_stocktracker_ui.py.",
        ],
    )

    _heading(doc, "11. Seguranca e boas praticas", 1)
    _bullet(
        doc,
        [
            "Nunca versionar config/secrets.py nem chaves reais no Git.",
            "Usar config/secrets.example.py apenas como modelo.",
            "data/stock.xlsx pode conter dados internos — tratar conforme politica Siemens.",
            "Nao executar modulos em src/gui/ isoladamente (imports relativos).",
        ],
    )

    _heading(doc, "12. Testes", 1)
    _table(
        doc,
        ["Tipo", "Como"],
        [
            ["GUI manual", "run.bat — fluxos SEARCH, SCAN, IN/OUT, historico"],
            ["Consola", "python -m src.test_terminal — menu 1–10"],
            ["API DigiKey", "python scripts/test_digikey_auth.py"],
        ],
    )
    _para(
        doc,
        "Registar em anexo: capturas de ecra, referencias testadas, erros e resolucoes.",
    )

    _heading(doc, "13. Limitacoes conhecidas", 1)
    _bullet(
        doc,
        [
            "Excel deve estar fechado ao gravar (PermissionError caso contrario).",
            "DigiKey sandbox: pesquisa pode falhar com 403 apesar de token valido.",
            "SEARCH nao consulta APIs — apenas Excel; catalogos usam SCAN.",
            "Nomes internos legados (ex.: coluna mouser no codigo = referencia fornecedor).",
        ],
    )

    _heading(doc, "14. Melhorias futuras", 1)
    _bullet(
        doc,
        [
            "Production App DigiKey para catalogo completo.",
            "Relatorios PDF / exportacao de stock baixo.",
            "Leitor de codigo de barras USB dedicado.",
            "Sincronizacao multi-utilizador (base de dados central).",
        ],
    )

    _heading(doc, "15. Conclusao", 1)
    _para(
        doc,
        "O Stock Tracker evoluiu de um prototipo em consola (v0.1) para uma aplicacao "
        "grafica Siemens completa (v2.0), mantendo sempre a separacao core/GUI. "
        "As versoes intermédias (GUI legada, demos com combo de distribuidor) "
        "documentam o percurso de desenvolvimento do estagio. A versao atual unifica "
        "funcionalidades e pesquisa multi-distribuidor, estando pronta para "
        "demonstracao e uso interno.",
    )

    _heading(doc, "Anexo A — Dependencias Python", 1)
    _bullet(doc, ["PySide6 >= 6.6", "openpyxl >= 3.1", "requests >= 2.31"])

    _heading(doc, "Anexo B — Documentacao no repositorio", 1)
    _bullet(
        doc,
        [
            "README.md — visao geral (EN)",
            "docs/user/COMMANDS.md — install and run",
            "docs/user/ARCHITECTURE.md — architecture",
            "docs/user/SUPPLIERS.md — APIs",
            "docs/user/DIGIKEY_SETUP.md — DigiKey setup",
            "docs/user/QT_DESIGNER.md — Qt Designer workflow",
            "src/gui/ESTRUTURA.md — mapa da GUI",
        ],
    )

    doc.save(OUT_FILE)
    return OUT_FILE


if __name__ == "__main__":
    path = build()
    print(f"Documento criado: {path}")
